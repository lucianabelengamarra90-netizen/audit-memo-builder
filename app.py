from flask import Flask, render_template, request, jsonify, send_file

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from docx import Document
from pypdf import PdfReader
from openai import OpenAI

from io import BytesIO, StringIO
from tempfile import NamedTemporaryFile
from datetime import datetime
from functools import lru_cache
from math import ceil

import csv
import os
import re
import sqlite3
import unicodedata
import uuid
import zipfile
import xml.etree.ElementTree as ET


app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 250 * 1024 * 1024


# ============================================================
# CONFIGURACIÓN
# ============================================================

ALLOWED_EXTENSIONS = {"xlsx", "csv", "docx", "pdf", "txt"}
MAX_ROW_TEXT_LENGTH = 12000
PROGRESS_LOG_EVERY_ROWS = 100000

HALLAZGO_PATTERN = re.compile(r"\bhallazgos?\b")


def normalize_text(value):
    if value is None:
        return ""
    text = str(value).strip()
    if not text:
        return ""
    text = unicodedata.normalize("NFD", text)
    text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
    text = re.sub(r"\s+", " ", text.lower())
    return text.strip()


def clean_text(value):
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value)).strip()


def get_extension(filename):
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def row_to_text(values):
    parts = [clean_text(v) for v in values if clean_text(v)]
    result = " | ".join(parts)
    return result[:MAX_ROW_TEXT_LENGTH]


def detect_categories(text):
    normalized = normalize_text(text)
    if not normalized:
        return []
    match = HALLAZGO_PATTERN.search(normalized)
    return [{"category": "Hallazgo", "keyword": match.group(0)}] if match else []


def column_letter_from_ref(cell_ref):
    if not cell_ref:
        return ""
    match = re.match(r"([A-Z]+)", cell_ref.upper())
    return match.group(1) if match else ""


def make_item(category, text, filename, origin_type, origin_name="", reference="", keyword=""):
    return {
        "id": str(uuid.uuid4()),
        "category": category,
        "text": clean_text(text),
        "filename": filename,
        "originType": origin_type,
        "originName": origin_name,
        "reference": reference,
        "keyword": keyword,
        "included": False,
        "converted": False,
    }


def add_unique_item(items, seen, item):
    fingerprint = (
        item.get("filename", ""), item.get("originName", ""),
        item.get("reference", ""), item.get("category", ""),
        normalize_text(item.get("text", ""))
    )
    if fingerprint in seen:
        return False
    seen.add(fingerprint)
    items.append(item)
    return True


def get_openai_client():
    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    return OpenAI(api_key=api_key) if api_key else None


MAIN_NS = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def qname(namespace, tag):
    return f"{{{namespace}}}{tag}"


def create_shared_strings_index(archive, db_path):
    connection = sqlite3.connect(db_path)
    cursor = connection.cursor()
    cursor.execute("PRAGMA journal_mode=OFF")
    cursor.execute("PRAGMA synchronous=OFF")
    cursor.execute("PRAGMA temp_store=FILE")
    cursor.execute("CREATE TABLE IF NOT EXISTS shared_strings (idx INTEGER PRIMARY KEY, value TEXT)")

    relevant_map = {}
    if "xl/sharedStrings.xml" not in archive.namelist():
        connection.commit()
        return connection, relevant_map, 0

    print("Analizando sharedStrings...", flush=True)
    batch = []
    index = 0

    with archive.open("xl/sharedStrings.xml") as stream:
        for _, element in ET.iterparse(stream, events=("end",)):
            if element.tag != qname(MAIN_NS, "si"):
                continue
            parts = []
            for text_element in element.iter(qname(MAIN_NS, "t")):
                if text_element.text:
                    parts.append(text_element.text)
            value = "".join(parts)
            matches = detect_categories(value)
            if matches:
                relevant_map[index] = {
                    "value": value,
                    "matches": matches,
                }
            batch.append((index, value))
            index += 1
            if len(batch) >= 10000:
                cursor.executemany("INSERT INTO shared_strings (idx, value) VALUES (?, ?)", batch)
                connection.commit()
                batch.clear()
            element.clear()

    if batch:
        cursor.executemany("INSERT INTO shared_strings (idx, value) VALUES (?, ?)", batch)
        connection.commit()

    print(
        f"SharedStrings analizados: {index} | índices relevantes: {len(relevant_map)}",
        flush=True
    )
    return connection, relevant_map, index


def get_workbook_sheet_paths(archive):
    workbook_root = ET.fromstring(archive.read("xl/workbook.xml"))
    relationships_root = ET.fromstring(archive.read("xl/_rels/workbook.xml.rels"))
    relationships = {}
    for relationship in relationships_root:
        rel_id = relationship.attrib.get("Id")
        target = relationship.attrib.get("Target", "")
        if not target:
            continue
        if target.startswith("/"):
            target = target.lstrip("/")
        elif not target.startswith("xl/"):
            target = "xl/" + target
        target = os.path.normpath(target).replace("\\", "/")
        relationships[rel_id] = target

    sheets = []
    sheets_element = workbook_root.find(qname(MAIN_NS, "sheets"))
    if sheets_element is None:
        return sheets
    archive_names = set(archive.namelist())
    for sheet in sheets_element:
        rel_id = sheet.attrib.get(qname(REL_NS, "id"))
        path = relationships.get(rel_id)
        if path and path in archive_names:
            sheets.append({
                "name": sheet.attrib.get("name", "Sin nombre"),
                "path": path,
                "state": sheet.attrib.get("state", "visible"),
            })
    return sheets


def extract_inline_string(cell):
    inline = cell.find(qname(MAIN_NS, "is"))
    if inline is None:
        return ""
    return "".join(
        text.text or ""
        for text in inline.iter(qname(MAIN_NS, "t"))
    )


def get_raw_cell_value(cell):
    value = cell.find(qname(MAIN_NS, "v"))
    return value.text if value is not None and value.text is not None else ""


def extract_xlsx(uploaded_file, filename):
    items, warnings, seen = [], [], set()
    excel_path = db_path = None
    shared_connection = None

    try:
        with NamedTemporaryFile(suffix=".xlsx", delete=False) as temp_excel:
            excel_path = temp_excel.name
            while True:
                chunk = uploaded_file.stream.read(1024 * 1024)
                if not chunk:
                    break
                temp_excel.write(chunk)

        with NamedTemporaryFile(suffix=".sqlite", delete=False) as temp_db:
            db_path = temp_db.name

        with zipfile.ZipFile(excel_path, "r") as archive:
            shared_connection, relevant_map, _ = create_shared_strings_index(archive, db_path)

            @lru_cache(maxsize=12000)
            def get_shared_value(index):
                row = shared_connection.execute(
                    "SELECT value FROM shared_strings WHERE idx = ?", (index,)
                ).fetchone()
                return row[0] if row else ""

            sheets = get_workbook_sheet_paths(archive)
            print(f"{filename}: {len(sheets)} solapas detectadas.", flush=True)

            for sheet_index, sheet in enumerate(sheets, start=1):
                sheet_name = sheet["name"]
                print(f"[{sheet_index}/{len(sheets)}] Procesando solapa: {sheet_name}", flush=True)

                rows_processed = 0
                relevant_rows = 0
                extracted_items = 0
                with archive.open(sheet["path"]) as stream:
                    for _, row_element in ET.iterparse(stream, events=("end",)):
                        if row_element.tag != qname(MAIN_NS, "row"):
                            continue

                        rows_processed += 1
                        if rows_processed % PROGRESS_LOG_EVERY_ROWS == 0:
                            print(
                                f"  {sheet_name}: {rows_processed:,} filas recorridas...",
                                flush=True
                            )

                        row_number = row_element.attrib.get("r", "")
                        descriptors = []
                        direct_matches = []
                        nonempty_cells = 0

                        for cell in row_element.findall(qname(MAIN_NS, "c")):
                            cell_type = cell.attrib.get("t")
                            cell_ref = cell.attrib.get("r", "")
                            column = column_letter_from_ref(cell_ref)

                            if cell_type == "s":
                                raw = get_raw_cell_value(cell)
                                if not raw:
                                    continue
                                try:
                                    shared_index = int(raw)
                                except ValueError:
                                    continue
                                descriptors.append(("s", shared_index))
                                nonempty_cells += 1

                                info = relevant_map.get(shared_index)
                                if info:
                                    direct_matches.extend(info.get("matches") or [])

                            elif cell_type == "inlineStr":
                                value = extract_inline_string(cell)
                                if not value:
                                    continue
                                descriptors.append(("text", value))
                                nonempty_cells += 1
                                matches = detect_categories(value)
                                direct_matches.extend(matches)

                            else:
                                raw = get_raw_cell_value(cell)
                                if not raw:
                                    continue
                                descriptors.append(("text", raw))
                                nonempty_cells += 1
                                if cell_type == "str":
                                    matches = detect_categories(raw)
                                    direct_matches.extend(matches)

                        if nonempty_cells == 0:
                            row_element.clear()
                            continue

                        if not direct_matches:
                            row_element.clear()
                            continue

                        values = []
                        for kind, value in descriptors:
                            values.append(get_shared_value(value) if kind == "s" else value)
                        row_text = row_to_text(values)
                        if not row_text:
                            row_element.clear()
                            continue

                        relevant_rows += 1
                        unique_matches = []
                        seen_match_keys = set()
                        for match in direct_matches:
                            key = (match.get("category", ""), match.get("keyword", ""))
                            if key not in seen_match_keys:
                                seen_match_keys.add(key)
                                unique_matches.append(match)

                        for match in unique_matches:
                            item = make_item(
                                match.get("category", "Información"),
                                row_text,
                                filename,
                                "Excel",
                                sheet_name,
                                f"Fila {row_number}" if row_number else "",
                                match.get("keyword", ""),
                            )
                            if add_unique_item(items, seen, item):
                                extracted_items += 1

                        row_element.clear()

                print(
                    f"Solapa finalizada: {sheet_name} | Filas: {rows_processed} | "
                    f"Filas relevantes: {relevant_rows} | Extraídas: {extracted_items}",
                    flush=True
                )

    except zipfile.BadZipFile:
        raise ValueError("El archivo Excel no es válido o está dañado.")
    finally:
        if shared_connection:
            try:
                shared_connection.close()
            except Exception:
                pass
        for path in (excel_path, db_path):
            if path:
                try:
                    os.remove(path)
                except Exception:
                    pass

    return items, warnings


def extract_csv(uploaded_file, filename):
    items, seen = [], set()
    raw = uploaded_file.read()
    try:
        text = raw.decode("utf-8-sig")
    except Exception:
        text = raw.decode("latin1", errors="replace")
    stream = StringIO(text)
    try:
        dialect = csv.Sniffer().sniff(text[:5000], delimiters=",;\t|")
    except Exception:
        dialect = csv.excel
    for row_number, row in enumerate(csv.reader(stream, dialect), start=1):
        row_text = row_to_text(row)
        if not row_text:
            continue
        for match in detect_categories(row_text):
            add_unique_item(items, seen, make_item(
                match["category"], row_text, filename, "CSV", "CSV",
                f"Fila {row_number}", match["keyword"]
            ))
    return items


def extract_docx(uploaded_file, filename):
    items, seen = [], set()
    document = Document(BytesIO(uploaded_file.read()))
    for number, paragraph in enumerate(document.paragraphs, start=1):
        text = clean_text(paragraph.text)
        if not text:
            continue
        for match in detect_categories(text):
            add_unique_item(items, seen, make_item(
                match["category"], text, filename, "Word", "Documento",
                f"Párrafo {number}", match["keyword"]
            ))
    for table_number, table in enumerate(document.tables, start=1):
        for row_number, row in enumerate(table.rows, start=1):
            text = row_to_text([cell.text for cell in row.cells])
            if not text:
                continue
            for match in detect_categories(text):
                add_unique_item(items, seen, make_item(
                    match["category"], text, filename, "Word", f"Tabla {table_number}",
                    f"Fila {row_number}", match["keyword"]
                ))
    return items


def extract_pdf(uploaded_file, filename):
    items, warnings, seen = [], [], set()
    reader = PdfReader(BytesIO(uploaded_file.read()))
    pages_with_text = 0
    for page_number, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
        if not page_text.strip():
            continue
        pages_with_text += 1
        lines = [clean_text(line) for line in page_text.splitlines() if clean_text(line)]
        for idx, line in enumerate(lines):
            matches = detect_categories(line)
            if not matches:
                continue
            context = " ".join(lines[max(0, idx - 1): min(len(lines), idx + 3)])
            for match in matches:
                add_unique_item(items, seen, make_item(
                    match["category"], context, filename, "PDF", f"Página {page_number}",
                    f"Página {page_number}", match["keyword"]
                ))
    if pages_with_text == 0:
        warnings.append(
            f"{filename}: no se pudo extraer texto del PDF; el documento puede estar escaneado."
        )
    return items, warnings


def extract_txt(uploaded_file, filename):
    items, seen = [], set()
    raw = uploaded_file.read()
    text = None
    for encoding in ("utf-8-sig", "utf-8", "latin1"):
        try:
            text = raw.decode(encoding)
            break
        except Exception:
            pass
    if text is None:
        return items
    for line_number, line in enumerate(text.splitlines(), start=1):
        line = clean_text(line)
        if not line:
            continue
        for match in detect_categories(line):
            add_unique_item(items, seen, make_item(
                match["category"], line, filename, "TXT", "Documento",
                f"Línea {line_number}", match["keyword"]
            ))
    return items


def extract_free_text(text):
    items, seen = [], set()
    for line_number, line in enumerate(text.splitlines(), start=1):
        line = clean_text(line)
        if not line:
            continue
        for match in detect_categories(line):
            add_unique_item(items, seen, make_item(
                match["category"], line, "Texto ingresado", "Texto", "Ingreso manual",
                f"Línea {line_number}", match["keyword"]
            ))
    return items


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/health")
def health():
    return jsonify({
        "status": "ok",
        "openaiConfigured": bool(os.getenv("OPENAI_API_KEY", "").strip())
    })


@app.route("/extract", methods=["POST"])
def extract_information():
    files = request.files.getlist("files")
    free_text = request.form.get("freeText", "").strip()
    if not files and not free_text:
        return jsonify({"error": "Cargá al menos un papel de trabajo o ingresá texto adicional."}), 400

    extracted_items, errors, warnings = [], [], []
    for uploaded_file in files:
        filename = uploaded_file.filename or "archivo_sin_nombre"
        extension = get_extension(filename)
        if extension not in ALLOWED_EXTENSIONS:
            errors.append(f"{filename}: formato no admitido.")
            continue
        print(f"Iniciando extracción: {filename}", flush=True)
        try:
            if extension == "xlsx":
                excel_items, excel_warnings = extract_xlsx(uploaded_file, filename)
                extracted_items.extend(excel_items)
                warnings.extend(excel_warnings)
            elif extension == "csv":
                extracted_items.extend(extract_csv(uploaded_file, filename))
            elif extension == "docx":
                extracted_items.extend(extract_docx(uploaded_file, filename))
            elif extension == "pdf":
                pdf_items, pdf_warnings = extract_pdf(uploaded_file, filename)
                extracted_items.extend(pdf_items)
                warnings.extend(pdf_warnings)
            elif extension == "txt":
                extracted_items.extend(extract_txt(uploaded_file, filename))
        except Exception as exc:
            print(f"Error procesando {filename}: {type(exc).__name__}: {exc}", flush=True)
            errors.append(f"{filename}: {type(exc).__name__}: {exc}")
        print(f"Extracción finalizada: {filename}", flush=True)

    if free_text:
        extracted_items.extend(extract_free_text(free_text))

    return jsonify({
        "items": extracted_items,
        "count": len(extracted_items),
        "errors": errors,
        "warnings": warnings,
        "message": (
            f"Se identificaron {len(extracted_items)} filas con la palabra hallazgo o hallazgos. "
            "Revisalos antes de incorporarlos al memo."
        )
    })


@app.route("/analyze", methods=["POST"])
def analyze_alias():
    return extract_information()


@app.route("/improve-text", methods=["POST"])
def improve_text():
    data = request.get_json(silent=True) or {}
    text = clean_text(data.get("text", ""))
    field_type = clean_text(data.get("fieldType", "Texto de auditoría"))
    if not text:
        return jsonify({"error": "Primero escribí una idea o texto para mejorar."}), 400

    openai_client = get_openai_client()
    if not openai_client:
        return jsonify({"error": "Render no está entregando OPENAI_API_KEY al proceso."}), 500

    instructions = """
Actuá como especialista en redacción de Auditoría Interna, con experiencia en retail mayorista.
Mejorá exclusivamente el texto aportado por el auditor. No realices la auditoría ni inventes hechos.
No inventes importes, cantidades, fechas, períodos, proveedores, clientes, personas, sistemas,
documentos, controles, muestras, resultados, hallazgos, evidencia, causas ni conclusiones.
Conservá el significado original. Usá lenguaje claro, técnico, profesional y corporativo.
OBJETIVO: profesionalizá el propósito sin inventar procedimientos.
ALCANCE: ordená únicamente la información suministrada.
SITUACIÓN OBSERVADA: describí objetivamente los hechos aportados.
RIESGO: mejorá la formulación sin inventar impactos específicos.
PROPUESTA DE MEJORA: fortalecé la redacción orientándola a controles, trazabilidad y eficiencia,
sin alterar la intención original.
Devolvé solamente el texto final.
"""
    prompt = f"TIPO DE CAMPO:\n{field_type}\n\nTEXTO DEL AUDITOR:\n{text}"

    try:
        response = openai_client.responses.create(
            model="gpt-5.6-luna",
            instructions=instructions,
            input=prompt,
        )
        improved = clean_text(response.output_text)
        if not improved:
            return jsonify({"error": "La IA no devolvió una redacción."}), 500
        return jsonify({"original": text, "improved": improved})
    except Exception as exc:
        print(f"Error OpenAI /improve-text: {type(exc).__name__}: {exc}", flush=True)
        return jsonify({
            "error": "La solicitud a OpenAI no pudo completarse. Revisá el log de Render."
        }), 500


NAVY = "17365D"
BLUE = "1F4E78"
MID_BLUE = "5B9BD5"
LIGHT_BLUE = "D9EAF7"
VERY_LIGHT_BLUE = "F3F7FB"
TEXT_COLOR = "1F2937"
MUTED = "667085"
BORDER_COLOR = "D0D7DE"
WHITE = "FFFFFF"
RED_FILL = "FDECEC"
RED_FONT = "B42318"
YELLOW_FILL = "FFF4CC"
YELLOW_FONT = "8A6200"
GREEN_FILL = "EAF6EA"
GREEN_FONT = "2E7D32"

THIN_BORDER = Border(
    left=Side(style="thin", color=BORDER_COLOR),
    right=Side(style="thin", color=BORDER_COLOR),
    top=Side(style="thin", color=BORDER_COLOR),
    bottom=Side(style="thin", color=BORDER_COLOR),
)


def set_cell(cell, value=None, *, bold=False, color=TEXT_COLOR, fill=None,
             size=10, align="left", valign="top", wrap=True, border=None):
    if value is not None:
        cell.value = value
    cell.font = Font(name="Aptos", size=size, bold=bold, color=color)
    if fill:
        cell.fill = PatternFill("solid", fgColor=fill)
    cell.alignment = Alignment(horizontal=align, vertical=valign, wrap_text=wrap)
    if border:
        cell.border = border


def merge_section_title(ws, row, title, start_col=1, end_col=8):
    ws.merge_cells(start_row=row, start_column=start_col, end_row=row, end_column=end_col)
    cell = ws.cell(row=row, column=start_col)
    set_cell(cell, title, bold=True, color=WHITE, fill=BLUE, size=11, valign="center")
    ws.row_dimensions[row].height = 24
    return row + 1


def write_merged_text(ws, row, text, start_col=1, end_col=8, min_height=34):
    ws.merge_cells(start_row=row, start_column=start_col, end_row=row, end_column=end_col)
    cell = ws.cell(row=row, column=start_col)
    set_cell(cell, text or "", color=TEXT_COLOR, fill=WHITE, size=10, border=THIN_BORDER)
    approx_lines = max(1, ceil(len(str(text or "")) / 110))
    ws.row_dimensions[row].height = max(min_height, min(120, approx_lines * 17))
    return row + 1


def included_extracted(extracted, categories=None):
    result = []
    for item in extracted:
        if not item.get("included"):
            continue
        if categories is not None and item.get("category") not in categories:
            continue
        result.append(item)
    return result


def unique_texts(items):
    seen = set()
    result = []
    for item in items:
        text = clean_text(item.get("text", ""))
        key = normalize_text(text)
        if text and key not in seen:
            seen.add(key)
            result.append(text)
    return result


def severity_style(severity):
    value = normalize_text(severity)
    if value == "alta":
        return RED_FILL, RED_FONT
    if value == "baja":
        return GREEN_FILL, GREEN_FONT
    return YELLOW_FILL, YELLOW_FONT


def style_sheet_base(ws):
    ws.sheet_view.showGridLines = False
    ws.freeze_panes = None


def build_memo_sheet(ws, general, findings, extracted):
    style_sheet_base(ws)
    ws.sheet_properties.tabColor = NAVY
    widths = {"A": 6, "B": 24, "C": 40, "D": 32, "E": 34, "F": 19, "G": 13, "H": 17}
    for col, width in widths.items():
        ws.column_dimensions[col].width = width

    ws.merge_cells("A1:H2")
    title = clean_text(general.get("title", "")) or "Memo de Auditoría"
    set_cell(ws["A1"], f"MEMO – {title.upper()}", bold=True, color=WHITE, fill=NAVY,
             size=16, align="center", valign="center")
    ws.row_dimensions[1].height = 24
    ws.row_dimensions[2].height = 24

    metadata = [
        ("Área", general.get("area", ""), "Período", general.get("period", "")),
        ("Proceso", general.get("process", ""), "Auditor", general.get("auditor", "")),
    ]
    row = 4
    for left_label, left_value, right_label, right_value in metadata:
        set_cell(ws.cell(row, 1), left_label, bold=True, color=NAVY, fill=LIGHT_BLUE, border=THIN_BORDER)
        ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=4)
        set_cell(ws.cell(row, 2), left_value or "", border=THIN_BORDER)
        set_cell(ws.cell(row, 5), right_label, bold=True, color=NAVY, fill=LIGHT_BLUE, border=THIN_BORDER)
        ws.merge_cells(start_row=row, start_column=6, end_row=row, end_column=8)
        set_cell(ws.cell(row, 6), right_value or "", border=THIN_BORDER)
        ws.row_dimensions[row].height = 23
        row += 1

    row += 1
    row = merge_section_title(ws, row, "Objetivo")
    row = write_merged_text(ws, row, general.get("objective", ""))
    row += 1

    if clean_text(general.get("scope", "")):
        row = merge_section_title(ws, row, "Alcance")
        row = write_merged_text(ws, row, general.get("scope", ""))
        row += 1

    tasks = unique_texts(included_extracted(extracted, {"Tarea realizada"}))
    row = merge_section_title(ws, row, "Trabajo realizado")
    if tasks:
        for idx, task in enumerate(tasks, start=1):
            set_cell(ws.cell(row, 1), f"{idx})", bold=True, color=NAVY, align="center", border=THIN_BORDER)
            ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=8)
            set_cell(ws.cell(row, 2), task, border=THIN_BORDER)
            ws.row_dimensions[row].height = max(30, min(100, ceil(len(task) / 100) * 17))
            row += 1
    else:
        row = write_merged_text(ws, row, "", min_height=28)
    row += 1

    result_categories = {"Conclusión", "Resultado", "Diferencia", "Observación", "Incumplimiento", "Pendiente", "Comentario"}
    results = unique_texts(included_extracted(extracted, result_categories))
    if results:
        row = merge_section_title(ws, row, "Resultados y observaciones relevantes")
        for text in results:
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=8)
            set_cell(ws.cell(row, 1), f"• {text}", border=THIN_BORDER, fill=VERY_LIGHT_BLUE)
            ws.row_dimensions[row].height = max(30, min(100, ceil(len(text) / 110) * 17))
            row += 1
        row += 1

    row = merge_section_title(ws, row, "Hallazgos")
    headers = ["N°", "Título", "Situación observada", "Riesgo", "Propuesta de mejora", "Área responsable", "Criticidad", "Estado"]
    for col_idx, header in enumerate(headers, start=1):
        set_cell(ws.cell(row, col_idx), header, bold=True, color=WHITE, fill=NAVY,
                 size=9, align="center", valign="center", border=THIN_BORDER)
    ws.row_dimensions[row].height = 32
    row += 1

    if not findings:
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=8)
        set_cell(ws.cell(row, 1), "Sin hallazgos incorporados.", color=MUTED, fill=VERY_LIGHT_BLUE, border=THIN_BORDER)
        row += 1
    else:
        for idx, finding in enumerate(findings, start=1):
            values = [
                idx,
                finding.get("title", ""),
                finding.get("situation", ""),
                finding.get("risk", ""),
                finding.get("proposal", ""),
                finding.get("responsibleArea", ""),
                finding.get("severity", ""),
                finding.get("status", ""),
            ]
            fill = WHITE if idx % 2 else VERY_LIGHT_BLUE
            max_len = 0
            for col_idx, value in enumerate(values, start=1):
                cell = ws.cell(row, col_idx)
                set_cell(cell, value, fill=fill, border=THIN_BORDER, size=9,
                         align="center" if col_idx in {1, 7, 8} else "left")
                max_len = max(max_len, len(str(value or "")))
            sev_fill, sev_font = severity_style(finding.get("severity", ""))
            set_cell(ws.cell(row, 7), finding.get("severity", ""), bold=True,
                     color=sev_font, fill=sev_fill, align="center", border=THIN_BORDER, size=9)
            ws.row_dimensions[row].height = max(42, min(120, ceil(max_len / 55) * 16))
            row += 1

    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.print_options.horizontalCentered = True
    ws.sheet_view.zoomScale = 85
    ws.print_area = f"A1:H{max(row, 1)}"


def build_detail_sheet(ws, headers, rows, widths, severity_column=None):
    style_sheet_base(ws)
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{max(1, len(rows) + 1)}"
    for col_idx, header in enumerate(headers, start=1):
        set_cell(ws.cell(1, col_idx), header, bold=True, color=WHITE, fill=NAVY,
                 size=9, align="center", valign="center", border=THIN_BORDER)
    ws.row_dimensions[1].height = 32
    for col_idx, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(col_idx)].width = width

    for row_idx, row_values in enumerate(rows, start=2):
        fill = WHITE if row_idx % 2 else VERY_LIGHT_BLUE
        max_len = 0
        for col_idx, value in enumerate(row_values, start=1):
            set_cell(ws.cell(row_idx, col_idx), value, fill=fill, border=THIN_BORDER, size=9)
            max_len = max(max_len, len(str(value or "")))
        if severity_column:
            severity = row_values[severity_column - 1] if len(row_values) >= severity_column else ""
            sev_fill, sev_font = severity_style(severity)
            set_cell(ws.cell(row_idx, severity_column), severity, bold=True, color=sev_font,
                     fill=sev_fill, align="center", border=THIN_BORDER, size=9)
        ws.row_dimensions[row_idx].height = max(24, min(100, ceil(max_len / 80) * 15))


def validate_source_integrity(findings, sources):
    source_names = {clean_text(source.get("name", "")) for source in sources if clean_text(source.get("name", ""))}
    if not source_names:
        return []
    foreign = []
    for finding in findings:
        source_file = clean_text(finding.get("sourceFile", ""))
        if source_file and source_file not in source_names:
            foreign.append(source_file)
    return sorted(set(foreign))


@app.route("/export-excel", methods=["POST"])
def export_excel():
    payload = request.get_json(silent=True) or {}
    memo = payload.get("memo", {})
    general = memo.get("general", {}) or {}
    findings = memo.get("findings", []) or []
    sources = memo.get("sources", []) or []
    extracted = memo.get("extracted", []) or []

    foreign_sources = validate_source_integrity(findings, sources)
    if foreign_sources:
        return jsonify({
            "error": "La exportación fue bloqueada porque hay hallazgos asociados a fuentes que no pertenecen a esta auditoría: "
                     + ", ".join(foreign_sources)
        }), 400

    workbook = Workbook()
    ws_memo = workbook.active
    ws_memo.title = "Memo"
    build_memo_sheet(ws_memo, general, findings, extracted)

    ws_findings = workbook.create_sheet("Hallazgos")
    ws_findings.sheet_properties.tabColor = "C55A11"
    finding_headers = [
        "N°", "Título", "Situación observada", "Riesgo", "Propuesta de mejora",
        "Área responsable", "Responsable plan", "Criticidad", "Estado",
        "Fecha compromiso", "Base cuantitativa", "Archivo origen", "Solapa / origen",
        "Evidencia / referencia", "Ticket", "Seguimiento"
    ]
    finding_rows = []
    for idx, finding in enumerate(findings, start=1):
        finding_rows.append([
            idx, finding.get("title", ""), finding.get("situation", ""), finding.get("risk", ""),
            finding.get("proposal", ""), finding.get("responsibleArea", ""), finding.get("actionOwner", ""),
            finding.get("severity", ""), finding.get("status", ""), finding.get("targetDate", ""),
            finding.get("quantitativeBasis", ""), finding.get("sourceFile", ""), finding.get("sourceLocation", ""),
            finding.get("evidence", ""), finding.get("ticket", ""), finding.get("followUp", "")
        ])
    build_detail_sheet(
        ws_findings, finding_headers, finding_rows,
        [6, 28, 45, 35, 38, 20, 20, 12, 16, 16, 20, 28, 22, 28, 16, 30],
        severity_column=8
    )

    ws_sources = workbook.create_sheet("Fuentes")
    ws_sources.sheet_properties.tabColor = MID_BLUE
    source_headers = ["Nombre", "Tipo", "Referencia", "Descripción"]
    source_rows = [[
        source.get("name", ""), source.get("type", ""), source.get("reference", ""), source.get("description", "")
    ] for source in sources]
    build_detail_sheet(ws_sources, source_headers, source_rows, [38, 14, 28, 55])

    ws_trace = workbook.create_sheet("Trazabilidad")
    ws_trace.sheet_properties.tabColor = "70AD47"
    trace_headers = ["Tipo detectado", "Contenido", "Archivo", "Origen / solapa", "Referencia", "Palabra clave", "Incluido", "Convertido"]
    trace_rows = [[
        item.get("category", ""), item.get("text", ""), item.get("filename", ""), item.get("originName", ""),
        item.get("reference", ""), item.get("keyword", ""), "Sí" if item.get("included") else "No",
        "Sí" if item.get("converted") else "No"
    ] for item in extracted]
    build_detail_sheet(ws_trace, trace_headers, trace_rows, [18, 65, 34, 28, 15, 20, 10, 11])

    output = BytesIO()
    workbook.save(output)
    output.seek(0)

    safe_title = re.sub(r"[^A-Za-z0-9ÁÉÍÓÚÜÑáéíóúüñ _-]+", "", clean_text(general.get("title", "")))
    safe_title = re.sub(r"\s+", "_", safe_title).strip("_") or "Audit_Memo"
    filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"

    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
